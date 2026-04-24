"""
Build Development_Journey.docx — Premium Edition
- Every STEP starts on a new page
- Figures embedded inline at the relevant step
- Dataset sample images shown in Step 1 (one per class)
- Rich styling: colored headings, shaded table headers, step banner boxes
"""
import os, sys, io
import copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE   = r"c:\Users\aakar\ProjectsDev\GlaucomaDet"
MD     = os.path.join(BASE, "Development_Journey.md")
FIGS   = os.path.join(BASE, "IEEE_Figures")
DS     = os.path.join(BASE, "yolo_dataset", "val")
OUT    = os.path.join(BASE, "Development_Journey.docx")

# ── Color Palette ──────────────────────────────────────────────────────────────
DARK_NAVY   = RGBColor(0x0D, 0x1B, 0x3E)
MID_BLUE    = RGBColor(0x1A, 0x5F, 0xA8)
ACCENT_TEAL = RGBColor(0x00, 0x8B, 0x8B)
STEP_ORANGE = RGBColor(0xD9, 0x6B, 0x00)
STEP_BG     = "FFF3E0"   # light orange fill for step banner
HDR_BG      = "1A5FA8"   # table header fill (blue)
ROW_ALT     = "EBF5FB"   # alternating row fill
BORDER_CLR  = "1A5FA8"

DOC = Document()

# ── Page Setup ─────────────────────────────────────────────────────────────────
sec = DOC.sections[0]
sec.page_width   = Inches(8.5)
sec.page_height  = Inches(11)
sec.left_margin  = sec.right_margin  = Inches(1.0)
sec.top_margin   = sec.bottom_margin = Inches(1.0)

# ── XML helpers ────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{side}')
        tag.set(qn('w:val'),   kwargs.get('val',   'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',    '6'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', BORDER_CLR))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def para_shading(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)

def add_paragraph_border(para, hex_color="1A5FA8", size="12"):
    pPr = para._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    size)
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), hex_color)
        pb.append(el)
    pPr.append(pb)

# ── Content helpers ─────────────────────────────────────────────────────────────
def add_heading(text, level):
    p   = DOC.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        run.font.size      = Pt(20)
        run.bold           = True
        run.font.color.rgb = DARK_NAVY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(8)
    elif level == 2:
        run.font.size      = Pt(14)
        run.bold           = True
        run.font.color.rgb = MID_BLUE
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
    else:
        run.font.size      = Pt(12)
        run.bold           = True
        run.font.color.rgb = ACCENT_TEAL
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(3)
    return p

def add_step_banner(text):
    """Full-width shaded banner for STEP headings."""
    p   = DOC.add_paragraph()
    para_shading(p, STEP_BG)
    add_paragraph_border(p, "D96B00", "12")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run(text)
    run.font.size      = Pt(18)
    run.bold           = True
    run.font.color.rgb = STEP_ORANGE
    return p

def add_body(text, bold=False, italic=False, color=None):
    p   = DOC.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold      = bold
    run.italic    = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p

def add_code(text):
    if not text.strip():
        return
    p   = DOC.add_paragraph()
    para_shading(p, "F4F4F4")
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_after  = Pt(2)

def add_table_from_rows(rows):
    import re
    if len(rows) < 2:
        return
    data = [r for r in rows if not all(re.match(r'^[-: ]+$', c.strip()) for c in r if c.strip())]
    if not data:
        return
    n_cols = max(len(r) for r in data)
    t = DOC.add_table(rows=len(data), cols=n_cols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(data):
        for ci in range(n_cols):
            cell_text = row[ci].strip() if ci < len(row) else ''
            cell = t.cell(ri, ci)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri == 0:
                set_cell_bg(cell, HDR_BG)
                run = cell.paragraphs[0].add_run(cell_text)
                run.bold           = True
                run.font.size      = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                bg = ROW_ALT if ri % 2 == 0 else "FFFFFF"
                set_cell_bg(cell, bg)
                run = cell.paragraphs[0].add_run(cell_text)
                run.font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    DOC.add_paragraph()

def insert_figure(fname, caption, width=Inches(5.8)):
    path = os.path.join(FIGS, fname)
    if not os.path.exists(path):
        print(f"  [SKIP] Figure not found: {path}")
        return
    p   = DOC.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    cap = DOC.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.italic       = True
        r.font.size    = Pt(9.5)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    DOC.add_paragraph()
    print(f"  [FIG] {fname}")

def insert_dataset_image(img_path, label, caption):
    if not os.path.exists(img_path):
        print(f"  [SKIP] Dataset image not found: {img_path}")
        return
    p   = DOC.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(1.9))
    cap = DOC.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.italic    = True
        r.font.size = Pt(9)

# ── COVER PAGE ──────────────────────────────────────────────────────────────────
DOC.add_paragraph()

title_p = DOC.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("Glaucoma Detection")
tr.font.size      = Pt(32)
tr.bold           = True
tr.font.color.rgb = DARK_NAVY

sub1 = DOC.add_paragraph()
sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
s1r = sub1.add_run("Complete Development Journey")
s1r.font.size      = Pt(18)
s1r.italic         = True
s1r.font.color.rgb = MID_BLUE

DOC.add_paragraph()

sub2 = DOC.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
s2r = sub2.add_run("From Zero to a 3-Model AI Diagnostic System")
s2r.font.size      = Pt(13)
s2r.font.color.rgb = ACCENT_TEAL

DOC.add_paragraph()
DOC.add_paragraph()

# Divider line (simulated with border)
div = DOC.add_paragraph()
add_paragraph_border(div, "1A5FA8", "8")

DOC.add_paragraph()

auth_p = DOC.add_paragraph()
auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
ar = auth_p.add_run("Aakarsh Shrey")
ar.font.size      = Pt(14)
ar.bold           = True
ar.font.color.rgb = DARK_NAVY

dept_p = DOC.add_paragraph()
dept_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = dept_p.add_run("Department of Computer Science & Engineering  |  April 2026")
dr.font.size = Pt(11)

DOC.add_paragraph()

# Quick-ref stats box on cover
stats = [
    ("Models Trained", "3"),
    ("Dataset Size",   "5,000 images"),
    ("Best Accuracy",  "89.4% (YOLOv11)"),
    ("Smallest Model", "0.4 MB (MambaOut)"),
    ("Total Steps",    "13"),
]
t = DOC.add_table(rows=1, cols=5)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, (label, val) in enumerate(stats):
    cell = t.cell(0, ci)
    set_cell_bg(cell, "EBF5FB")
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(val + "\n")
    r1.bold           = True
    r1.font.size      = Pt(13)
    r1.font.color.rgb = MID_BLUE
    r2 = p1.add_run(label)
    r2.font.size      = Pt(8)
    r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

DOC.add_page_break()

# ── PARSE MARKDOWN AND RENDER ───────────────────────────────────────────────────
import re

with open(MD, encoding='utf-8') as f:
    lines = f.readlines()

# Map: which figure to insert AFTER a given STEP heading line
# key = text that appears in the heading, value = (fname, caption)
STEP_FIGS = {
    "STEP 3":  ("Fig1_Accuracy_Sensitivity.png", "Fig. 1 — YOLOv11 Baseline: Accuracy & Sensitivity by Class"),
    "STEP 8":  ("Fig3_Training_Curves.png",       "Fig. 3 — Training Convergence: Accuracy (a) and Loss (b) over 30 Epochs"),
    "STEP 9":  ("Fig4_Confusion_Matrices.png",    "Fig. 4 — Full Confusion Matrices for all Three Models (n=200/class)"),
    "STEP 10": ("Fig5_ROC_Curves.png",            "Fig. 5 — ROC/AUC Curves: Glaucoma Class One-vs-Rest"),
    "STEP 11": ("Fig2_Efficiency_Scatter.png",    "Fig. 2 — Computational Efficiency vs. Accuracy Trade-off"),
    "STEP 12": ("Fig6_Radar_Chart.png",           "Fig. 6 — Multi-Metric Radar Comparison of All Three Models"),
    "STEP 13": ("Fig7_Threshold_Analysis.png",    "Fig. 7 — Clinical Decision Threshold Analysis"),
}

# Dataset images for Step 1
DS_IMAGES = [
    (os.path.join(DS, "glaucoma",         "img_0.jpg"),    "Glaucoma",         "Sample — Glaucoma\n(Lower blue channel, reduced contrast)"),
    (os.path.join(DS, "glaucoma_suspect", "img_1234.jpg"), "Glaucoma Suspect", "Sample — Glaucoma Suspect\n(Borderline optic disc changes)"),
    (os.path.join(DS, "non_glaucoma",     "img_10.jpg"),   "Non-Glaucoma",     "Sample — Normal Retina\n(Healthy disc, clear margins)"),
]

in_code  = False
in_table = False
code_buf = []
tbl_rows = []
i = 0
pending_figure = None  # figure to insert after current block settles

STEP_HEADING_KEYS = list(STEP_FIGS.keys())

def is_step_heading(text):
    for k in STEP_HEADING_KEYS:
        if k in text:
            return k
    return None

while i < len(lines):
    raw  = lines[i]
    line = raw.rstrip('\n')

    # ── Code block toggle ────────────────────────────────────────
    if line.strip().startswith('```'):
        if in_code:
            add_code('\n'.join(code_buf))
            code_buf = []
            in_code  = False
        else:
            in_code = True
        i += 1
        continue

    if in_code:
        code_buf.append(line)
        i += 1
        continue

    # ── Table detection ──────────────────────────────────────────
    if '|' in line and line.strip().startswith('|'):
        cells = [c for c in line.strip().split('|') if c != '']
        tbl_rows.append(cells)
        i += 1
        next_line = lines[i].strip() if i < len(lines) else ''
        if not next_line.startswith('|'):
            add_table_from_rows(tbl_rows)
            tbl_rows = []
        continue

    if tbl_rows:
        add_table_from_rows(tbl_rows)
        tbl_rows = []

    # ── Headings ─────────────────────────────────────────────────
    if line.startswith('# '):
        txt = line[2:]
        step_key = is_step_heading(txt)
        if step_key or txt.strip().startswith('STEP') or txt.strip().startswith('Final') or txt.strip().startswith('Key'):
            DOC.add_page_break()
            add_step_banner(txt)
            if step_key and step_key in STEP_FIGS:
                pending_figure = STEP_FIGS[step_key]
        else:
            add_heading(txt, 1)

        # Dataset images after Step 1 heading
        if "STEP 1" in txt:
            DOC.add_paragraph()
            add_heading("Dataset Sample Images — One Per Class", 2)
            cap_p = DOC.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cap_p.add_run("Retinal fundus photographs from the Hugging Face SMDG dataset used for training.")
            cr.italic    = True
            cr.font.size = Pt(9)
            # Insert 3 images side by side via a 1-row, 3-col table
            img_tbl = DOC.add_table(rows=2, cols=3)
            img_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ci, (img_path, label, cap_text) in enumerate(DS_IMAGES):
                cell = img_tbl.cell(0, ci)
                set_cell_bg(cell, "FAFAFA")
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if os.path.exists(img_path):
                    run = cell.paragraphs[0].add_run()
                    run.add_picture(img_path, width=Inches(1.7))
                    print(f"  [DS IMG] {img_path}")
                else:
                    cell.paragraphs[0].add_run(f"[Image: {label}]")
                cap_cell = img_tbl.cell(1, ci)
                set_cell_bg(cap_cell, "EBF5FB")
                cap_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr2 = cap_cell.paragraphs[0].add_run(cap_text)
                cr2.italic    = True
                cr2.font.size = Pt(8.5)
                cr2.font.color.rgb = RGBColor(0x33, 0x33, 0x55)
            DOC.add_paragraph()

    elif line.startswith('## '):
        add_heading(line[3:], 2)
    elif line.startswith('### '):
        add_heading(line[4:], 3)
    elif line.strip() == '---':
        # Insert pending figure before the divider if one is queued
        if pending_figure:
            insert_figure(pending_figure[0], pending_figure[1])
            pending_figure = None
        hr = DOC.add_paragraph()
        add_paragraph_border(hr, "AAAAAA", "4")
    elif line.strip() == '':
        DOC.add_paragraph()
    else:
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
        clean = re.sub(r'\*(.+?)\*',     r'\1', clean)
        clean = re.sub(r'`(.+?)`',       r'\1', clean)
        clean = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', clean)
        clean = clean.lstrip('> ')
        if clean.strip():
            bold_line = line.strip().startswith('**') and line.strip().endswith('**')
            add_body(clean, bold=bold_line)

    i += 1

# flush any remaining pending figure
if pending_figure:
    insert_figure(pending_figure[0], pending_figure[1])

# ── SAVE ───────────────────────────────────────────────────────────────────────
DOC.save(OUT)
print(f"Done: {OUT}  ({os.path.getsize(OUT)//1024} KB)")
